"""
File Processing Module for Document Text Extraction

This module provides a comprehensive file processing system for extracting text content
from various document formats including PDF, DOCX, and plain text files. It features
robust error handling, detailed logging, and validation for production use in
document processing pipelines.

Author: SomuTech
Date: July 2025
License: MIT
"""

from __future__ import annotations

import logging
from typing import BinaryIO, Callable, Dict, Optional

import PyPDF2
from docx import Document

# Logging configuration
logger = logging.getLogger(__name__)

# Helpers for each file type


def _extract_pdf(file_obj: BinaryIO) -> str:
    """Extract text from a PDF file object."""
    reader = PyPDF2.PdfReader(file_obj)
    text = "".join((page.extract_text() or "") + "\n" for page in reader.pages)
    logger.info("PDF: extracted %s pages", len(reader.pages))
    return text


def _extract_docx(file_obj: BinaryIO) -> str:
    """Extract text from a DOCX file object."""
    doc = Document(file_obj)
    text = "".join(p.text + "\n" for p in doc.paragraphs)
    logger.info("DOCX: extracted %s paragraphs", len(doc.paragraphs))
    return text


def _extract_txt(file_obj: BinaryIO) -> str:
    """Extract text from a plain-text file object (UTF-8 with latin-1 fallback)."""
    try:
        content = file_obj.read().decode("utf-8")
        encoding = "utf-8"
    except UnicodeDecodeError:
        content = file_obj.read().decode("latin-1", errors="ignore")
        encoding = "latin-1"
    logger.info("TXT: extracted using %s decoding", encoding)
    return content


# Dispatcher mapping mime-types → extractor
_EXTRACTORS: Dict[str, Callable[[BinaryIO], str]] = {
    "application/pdf": _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
    "text/plain": _extract_txt,
}


# Public function
def extract_text_from_file(uploaded_file: BinaryIO) -> Optional[str]:
    """
    Extract raw text from an uploaded file.

    Parameters
    ----------
    uploaded_file : BinaryIO
        Streamlit-style uploaded file object having `.type`, `.read()` and/or iterator
        behaviour.

    Returns
    -------
    Optional[str]
        Extracted text or ``None`` if the type is unsupported or an error occurred.
    """
    mime_type: str = getattr(uploaded_file, "type", "")
    extractor = _EXTRACTORS.get(mime_type)

    if extractor is None:
        logger.warning("Unsupported file type: %s", mime_type)
        return None

    # Reset pointer to start for safety
    try:
        uploaded_file.seek(0)
    except Exception:  # file-like objects without seek()
        pass

    try:
        text = extractor(uploaded_file)
        return text if text.strip() else None
    except Exception as exc:  # noqa: BLE001
        logger.error("Failed to extract text (%s): %s", mime_type, exc, exc_info=True)
        return None
